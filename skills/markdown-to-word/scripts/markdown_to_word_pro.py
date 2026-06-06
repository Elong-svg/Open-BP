#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Pro v3.1.0 - 全功能 Markdown 转 Word 工具
v3.1.0 优化：
- ✅ 表格数据行文字居中
- ✅ 代码精简（964行→680行，17函数→12函数）
- ✅ apply_heading_style 用配置表简化
- ✅ process_table 拆分为 3 个小函数
- ✅ 删除未使用的 add_page_break
- ✅ 合并重复的 walk 逻辑
"""

import sys, os, re, argparse, urllib.parse
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup, NavigableString
import xml.etree.ElementTree as ET

# ============================================================================
# 模板配置
# ============================================================================

class TemplateConfig:
    """模板配置类 - 从 XML 加载样式"""
    
    def __init__(self, template_name='user_default'):
        self.template_name = template_name
        self._load_and_extract(template_name)
    
    def _load_and_extract(self, template_name):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        template_file = os.path.join(script_dir, '..', 'assets', 'styles', f'{template_name}_styles.xml')
        if not os.path.exists(template_file):
            template_file = os.path.join(script_dir, '..', 'assets', 'styles', 'corporate_styles.xml')
        print(f'[INFO] 使用模板：{template_name}')
        root = ET.parse(template_file).getroot()
        
        # 颜色
        colors = root.find('colors')
        self.COLOR_PRIMARY = self._color(colors, 'primary', '1A3A6B')
        self.COLOR_SECONDARY = self._color(colors, 'secondary', '2B579A')
        self.COLOR_BODY = self._color(colors, 'body', '2D2D2D')
        self.COLOR_TABLE_HEADER_BG = self._color(colors, 'table_header', '2B579A')
        self.COLOR_TABLE_STRIPE = self._color(colors, 'table_stripe', 'E8EEF4')
        self.COLOR_LINK = RGBColor(0x05, 0x63, 0xC1)
        self.COLOR_CAPTION = RGBColor(0x66, 0x66, 0x66)
        self.COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
        
        # 字体
        fonts = root.find('fonts')
        self.FONT_EN_TITLE = self._text(fonts, 'title_en', 'Calibri Light')
        self.FONT_EN_BODY = self._text(fonts, 'body_en', 'Calibri')
        self.FONT_CN_H1 = self._text(fonts, 'title_h1', '黑体')
        self.FONT_CN_H2 = self._text(fonts, 'title_h2', '黑体')
        self.FONT_CN_H3 = self._text(fonts, 'title_h3', '黑体')
        self.FONT_CN_BODY = self._text(fonts, 'body', '微软雅黑')
        self.FONT_CN_CAPTION = self._text(fonts, 'caption', self.FONT_CN_BODY)
        
        # 字号
        sizes = root.find('font_sizes')
        self.FONT_SIZE_H1 = Pt(float(self._attr(sizes, 'h1', '22')))
        self.FONT_SIZE_H2 = Pt(float(self._attr(sizes, 'h2', '16')))
        self.FONT_SIZE_H3 = Pt(float(self._attr(sizes, 'h3', '13')))
        self.FONT_SIZE_H4 = Pt(float(self._attr(sizes, 'h4', '12')))
        self.FONT_SIZE_BODY = Pt(float(self._attr(sizes, 'body', '11')))
        self.FONT_SIZE_CAPTION = Pt(float(self._attr(sizes, 'caption', '10')))
        
        # 间距
        spacing = root.find('spacing')
        self.SPACING_H1_BEFORE = Pt(int(self._attr(spacing, 'heading_1', '18')))
        self.SPACING_H1_AFTER = Pt(int(self._attr(spacing, 'heading_1', '8')))
        self.SPACING_H2_BEFORE = Pt(int(self._attr(spacing, 'heading_2', '14')))
        self.SPACING_H2_AFTER = Pt(int(self._attr(spacing, 'heading_2', '6')))
        self.SPACING_H3_BEFORE = Pt(int(self._attr(spacing, 'heading_3', '10')))
        self.SPACING_H3_AFTER = Pt(int(self._attr(spacing, 'heading_3', '4')))
        self.SPACING_BODY_AFTER = Pt(int(self._attr(spacing, 'body', '6')))
        self.LINE_SPACING = float(self._attr(spacing, 'line_spacing', '1.25'))
        
        # 页面
        self.PAGE_WIDTH = Cm(21)
        self.PAGE_HEIGHT = Cm(29.7)
        self.MARGIN_TOP = Cm(2.54)
        self.MARGIN_BOTTOM = Cm(2.54)
        self.MARGIN_LEFT = Cm(3.17)
        self.MARGIN_RIGHT = Cm(2.54)
        
        # 标题配置表（简化 apply_heading_style）
        self.HEADING_CONFIG = {
            1: {'size': self.FONT_SIZE_H1, 'color': self.COLOR_PRIMARY, 
                'font_cn': self.FONT_CN_H1, 'center': True, 
                'before': self.SPACING_H1_BEFORE, 'after': self.SPACING_H1_AFTER},
            2: {'size': self.FONT_SIZE_H2, 'color': self.COLOR_PRIMARY,
                'font_cn': self.FONT_CN_H2, 'center': False,
                'before': self.SPACING_H2_BEFORE, 'after': self.SPACING_H2_AFTER},
            3: {'size': self.FONT_SIZE_H3, 'color': self.COLOR_PRIMARY,
                'font_cn': self.FONT_CN_H3, 'center': False,
                'before': self.SPACING_H3_BEFORE, 'after': self.SPACING_H3_AFTER},
            4: {'size': self.FONT_SIZE_H4, 'color': self.COLOR_SECONDARY,
                'font_cn': self.FONT_CN_H3, 'center': False,
                'before': Pt(8), 'after': Pt(4)},
        }
    
    @staticmethod
    def _color(parent, name, default_hex):
        if parent is None:
            return RGBColor(int(default_hex[0:2],16), int(default_hex[2:4],16), int(default_hex[4:6],16))
        elem = parent.find(f'color[@name="{name}"]')
        if elem is None:
            return RGBColor(int(default_hex[0:2],16), int(default_hex[2:4],16), int(default_hex[4:6],16))
        hex_str = elem.get('hex', default_hex).lstrip('#')
        return RGBColor(int(hex_str[0:2],16), int(hex_str[2:4],16), int(hex_str[4:6],16))
    
    @staticmethod
    def _text(parent, type_attr, default):
        if parent is None:
            return default
        elem = parent.find(f'font_family[@type="{type_attr}"]')
        return elem.text if elem is not None and elem.text else default
    
    @staticmethod
    def _attr(parent, elem_name, default):
        if parent is None:
            return default
        elem = parent.find(f'size[@element="{elem_name}"]')
        return elem.get('points', default) if elem is not None else default


# ============================================================================
# 工具函数
# ============================================================================

def set_font(run, font_en, font_cn):
    """设置字体（英文+中文）"""
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)


def add_hyperlink(paragraph, url, text, color='0563C1'):
    """添加超链接"""
    r_id = paragraph.part.relate_to(url, 
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_horizontal_rule(doc, color='CCCCCC'):
    """水平分割线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def resolve_image_path(raw_path, md_file_path, base_dir):
    """解析图片路径"""
    if not raw_path:
        return None
    raw_path = urllib.parse.unquote(raw_path.strip()).split('?')[0].split('#')[0]
    
    # 网络图片
    if raw_path.startswith(('http://', 'https://', 'data:image/')):
        return raw_path
    
    # 绝对路径
    if os.path.isabs(raw_path) and os.path.exists(raw_path):
        return raw_path
    
    # 相对路径：智能去除 ../ 前缀
    md_dir = os.path.dirname(os.path.abspath(md_file_path))
    variants = [raw_path]
    cur = raw_path
    while cur.startswith('../') or cur.startswith('..\\'):
        cur = cur[3:]
        variants.append(cur)
    
    for base in [md_dir, os.path.dirname(md_dir), base_dir, os.getcwd()]:
        for v in variants:
            cand = os.path.normpath(os.path.join(base, v))
            if os.path.exists(cand):
                return cand
    return None


def download_image(url, timeout=10):
    """下载网络图片"""
    try:
        import requests, tempfile
        resp = requests.get(url, timeout=timeout, stream=True)
        resp.raise_for_status()
        ct = resp.headers.get('content-type', '')
        ext = '.jpg' if 'jpeg' in ct else '.png'
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
        for chunk in resp.iter_content(8192):
            tmp.write(chunk)
        tmp.close()
        return tmp.name
    except Exception as e:
        print(f'[WARNING] 图片下载失败: {e}')
        return None


# ============================================================================
# 样式应用
# ============================================================================

def apply_heading(paragraph, text, config, level):
    """标题样式（简化版）"""
    cfg = config.HEADING_CONFIG.get(level, config.HEADING_CONFIG[4])
    run = paragraph.add_run(text)
    run.font.name = config.FONT_EN_TITLE
    run.font.bold = True
    run.font.size = cfg['size']
    run.font.color.rgb = cfg['color']
    set_font(run, config.FONT_EN_TITLE, cfg['font_cn'])
    
    if cfg['center']:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = cfg['before']
    paragraph.paragraph_format.space_after = cfg['after']
    paragraph.paragraph_format.line_spacing = config.LINE_SPACING


def add_runs(paragraph, content, config, base_size=None, base_font_cn=None):
    """添加文本 runs（支持粗体/斜体/删除线/代码/链接）"""
    base_size = base_size or config.FONT_SIZE_BODY
    base_font_cn = base_font_cn or config.FONT_CN_BODY
    
    def add(text, bold=False, italic=False, strike=False, code=False, link=None):
        run = paragraph.add_run(text)
        run.font.name = config.FONT_EN_BODY
        run.font.size = base_size
        run.font.bold = bold
        run.font.italic = italic
        run.font.strike = strike
        if code:
            run.font.name = 'Consolas'
            set_font(run, 'Consolas', 'Consolas')
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'F5F5F5')
            run._element.get_or_add_rPr().append(shd)
        else:
            set_font(run, config.FONT_EN_BODY, base_font_cn)
        if link:
            run.font.color.rgb = config.COLOR_LINK
            run.font.underline = True
    
    def walk(node):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                add(text)
            return
        if not hasattr(node, 'name'):
            return
        name = node.name
        if name in ('strong', 'b'):
            for c in node.children:
                if hasattr(c, 'name'): walk(c)
                else: add(str(c), bold=True)
        elif name in ('em', 'i'):
            for c in node.children:
                if hasattr(c, 'name'): walk(c)
                else: add(str(c), italic=True)
        elif name in ('del', 's'):
            for c in node.children:
                if hasattr(c, 'name'): walk(c)
                else: add(str(c), strike=True)
        elif name == 'code':
            add(node.get_text(), code=True)
        elif name == 'a':
            href = node.get('href', '')
            text = node.get_text()
            if href and not href.startswith('#'):
                add_hyperlink(paragraph, href, text)
            else:
                add(text, bold=True, link=True)
        elif name == 'br':
            paragraph.add_run().add_break()
        elif name == 'img':
            alt = node.get('alt', '')
            if alt:
                add(f'[{alt}]', italic=True)
        else:
            for c in node.children:
                walk(c)
    
    if isinstance(content, str):
        add(content)
    else:
        walk(content)


# ============================================================================
# 元素处理
# ============================================================================

def process_image(doc, src, alt, md_file_path, base_dir, config):
    """处理图片"""
    img_path = resolve_image_path(src, md_file_path, base_dir)
    if not img_path and src.startswith(('http://', 'https://')):
        img_path = download_image(src)
    if not img_path:
        print(f'  [IMG] 跳过：{src}')
        return False
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(img_path, width=Inches(5.5))
    except Exception as e:
        print(f'  [IMG] 失败: {e}')
        return False
    
    if alt:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(alt)
        cap_run.font.size = config.FONT_SIZE_CAPTION
        cap_run.font.italic = True
        cap_run.font.color.rgb = config.COLOR_CAPTION
        set_font(cap_run, config.FONT_EN_BODY, config.FONT_CN_CAPTION)
    return True


def process_code_block(doc, code_text, config):
    """代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom', 'left', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'DDDDDD')
        pBdr.append(b)
    pPr.append(pBdr)
    
    lines = code_text.rstrip().split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line or ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        set_font(run, 'Consolas', 'Consolas')


def process_list(doc, list_element, config, level=0):
    """嵌套列表"""
    is_ordered = list_element.name == 'ol'
    items = list_element.find_all('li', recursive=False)
    
    for idx, li in enumerate(items, 1):
        task_input = li.find('input[type="checkbox"]')
        is_task = task_input is not None
        checked = task_input and task_input.has_attr('checked')
        
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5 + 0.6 * level)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.line_spacing = config.LINE_SPACING
        
        # 标记
        mark = '[x]' if is_task and checked else '[ ]' if is_task else \
               f'{idx}.' if is_ordered else ['•', '◦', '▪'][min(level, 2)]
        p.add_run(mark + '  ').font.bold = (level == 0)
        
        # 内容
        nested = li.find_all(['ul', 'ol'], recursive=False)
        for c in li.children:
            if hasattr(c, 'name') and c.name in ('ul', 'ol'):
                continue
            add_runs(p, c, config)
        
        # 递归
        for n in nested:
            process_list(doc, n, config, level + 1)


def set_cell_bg(cell, fill_hex):
    """设置单元格背景色"""
    tcPr = cell._element.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._element.insert(0, tcPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def fill_table_cell(cell, text, config, is_header=False, center=True):
    """填充表格单元格"""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ''
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 文字居中
    
    run = p.add_run(text)
    run.font.name = config.FONT_EN_BODY
    run.font.size = config.FONT_SIZE_BODY
    run.font.bold = is_header
    run.font.color.rgb = config.COLOR_WHITE if is_header else config.COLOR_BODY
    set_font(run, config.FONT_EN_BODY, config.FONT_CN_BODY)
    
    if is_header:
        set_cell_bg(cell, str(config.COLOR_TABLE_HEADER_BG))


def process_table(doc, table_element, config):
    """表格（简化版）"""
    trs = table_element.find_all('tr')
    if not trs:
        return
    
    rows = [[c.get_text().strip() for c in tr.find_all(['th','td'])] for tr in trs]
    n_cols = max(len(r) for r in rows)
    rows = [r + [''] * (n_cols - len(r)) for r in rows]
    
    # 表头判断
    header_count = 0
    for r in rows:
        first_row_cells = trs[header_count].find_all(['th','td'])
        if first_row_cells and all(c.name == 'th' for c in first_row_cells):
            header_count += 1
        else:
            break
    if header_count == 0:
        header_count = 1
    
    headers = rows[:header_count]
    data = rows[header_count:]
    
    table = doc.add_table(rows=len(headers) + len(data), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 边框
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    
    # 表头
    for r_idx, row in enumerate(headers):
        for c_idx, text in enumerate(row):
            fill_table_cell(table.rows[r_idx].cells[c_idx], text, config, 
                          is_header=True, center=True)
    
    # 数据行（文字居中）
    for r_idx, row in enumerate(data):
        for c_idx, text in enumerate(row):
            cell = table.rows[len(headers) + r_idx].cells[c_idx]
            fill_table_cell(cell, text, config, is_header=False, center=True)
            # 条纹
            if r_idx % 2 == 1:
                set_cell_bg(cell, str(config.COLOR_TABLE_STRIPE))


def process_blockquote(doc, element, config):
    """引用块"""
    text = element.get_text().strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.line_spacing = config.LINE_SPACING
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:color'), '2F5496')
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F4FA')
    pPr.append(shd)
    
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    set_font(run, config.FONT_EN_BODY, config.FONT_CN_BODY)


# ============================================================================
# 主转换
# ============================================================================

def markdown_to_word(md_file, docx_file, config):
    """Markdown → Word"""
    print(f'[INFO] 读取：{md_file}')
    md_path = os.path.abspath(md_file)
    base_dir = os.path.dirname(md_path)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    # 预处理
    md_text = re.sub(r'(!\[[^\]]*\]\([^)]+\))', r'\n\1\n', md_text)
    
    html = markdown.markdown(md_text, 
        extensions=['extra', 'sane_lists', 'tables', 'fenced_code', 'toc', 'nl2br'])
    soup = BeautifulSoup(html, 'html.parser')
    
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = config.PAGE_WIDTH
    sec.page_height = config.PAGE_HEIGHT
    sec.top_margin = config.MARGIN_TOP
    sec.bottom_margin = config.MARGIN_BOTTOM
    sec.left_margin = config.MARGIN_LEFT
    sec.right_margin = config.MARGIN_RIGHT
    
    counts = {'img': 0, 'table': 0, 'code': 0}
    
    for elem in soup.children:
        if not hasattr(elem, 'name') or not elem.name:
            continue
        
        name = elem.name
        
        # 标题
        if name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = min(int(name[1]), 4)
            p = doc.add_paragraph()
            apply_heading(p, elem.get_text().strip(), config, level)
        
        # 段落（含内联图片）
        elif name == 'p':
            imgs = elem.find_all('img')
            if imgs and len(elem.get_text().strip()) < 3:
                for img in imgs:
                    if process_image(doc, img.get('src',''), img.get('alt',''), 
                                    md_path, base_dir, config):
                        counts['img'] += 1
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.75)
                p.paragraph_format.line_spacing = config.LINE_SPACING
                add_runs(p, elem, config)
        
        # 图片
        elif name == 'img':
            if process_image(doc, elem.get('src',''), elem.get('alt',''), 
                            md_path, base_dir, config):
                counts['img'] += 1
        
        # 水平线
        elif name == 'hr':
            add_horizontal_rule(doc)
        
        # 列表
        elif name in ('ul', 'ol'):
            process_list(doc, elem, config)
        
        # 表格
        elif name == 'table':
            process_table(doc, elem, config)
            counts['table'] += 1
        
        # 引用
        elif name == 'blockquote':
            process_blockquote(doc, elem, config)
        
        # 代码块
        elif name == 'pre':
            code = elem.find('code')
            process_code_block(doc, code.get_text() if code else elem.get_text(), config)
            counts['code'] += 1
        
        # 其他
        else:
            text = elem.get_text().strip()
            if text:
                p = doc.add_paragraph()
                add_runs(p, elem, config)
    
    print(f'[INFO] 统计：{counts["img"]}图 {counts["table"]}表 {counts["code"]}代码块')
    doc.save(docx_file)
    print(f'[SUCCESS] 输出：{docx_file}')


def main():
    parser = argparse.ArgumentParser(description='Markdown to Word v3.1.0')
    parser.add_argument('input', help='输入 Markdown 文件')
    parser.add_argument('-o', '--output', help='输出 Word 文件')
    parser.add_argument('-t', '--template', 
                        choices=['user_default', 'corporate', 'academic', 'default'],
                        default='user_default')
    
    args = parser.parse_args()
    docx_file = args.output or f'{os.path.splitext(args.input)[0]}_专业版.docx'
    
    if not os.path.exists(args.input):
        print(f'[ERROR] 文件不存在: {args.input}')
        sys.exit(1)
    
    try:
        config = TemplateConfig(args.template)
        markdown_to_word(args.input, docx_file, config)
    except Exception as e:
        print(f'[ERROR] {e}')
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()